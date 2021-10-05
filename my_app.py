from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# users profile picture
document.add_picture('C:\\Users\\alexa\PycharmProjects\cv-project\pexels-cottonbro-5474040.jpg', width=Inches(2.0))

# taking users details
name = input('What is your name? ').capitalize()
speak(f'Hello {name} how are you doing today')
speak('What is your number? ')
phone_number = input('What is your number? ')
email = input('What is your email address? ')

# paragraphing users input
document.add_paragraph(f'{name} | {phone_number} | {email}')


def work_experience():
    # work experience
    p = document.add_paragraph()

    company = input('Name of the company? ').title()
    from_date = input('From date? ')
    to_date = input('To date? ')

    # adding words on the same paragraph
    p.add_run(f'{company}').bold = True
    p.add_run(f' {from_date} - {to_date} \n').italic = True

    experience_details = input(f'Describe your experience at {company}? ')
    p.add_run(experience_details)


# about the user
document.add_heading('About Me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)
document.add_heading('Work Experience')
work_experience()

# more experiences
while True:
    has_more_experience = input('Do you have more experiences? (Yes/No )').lower()
    if has_more_experience == 'yes':
        work_experience()
    else:
        break

# users skills
document.add_heading('Skills')


def skills():
    skill = input('What is skills do you have? ').capitalize()
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'


skills()
# more skills
while True:
    more_skills = input('Do you have more skills? (yes/no) ').lower()
    if more_skills == 'yes':
        skills()
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Alexanders cv generator "
document.save('cv2.docx')
